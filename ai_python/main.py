import os
import uuid
import threading
import io
import base64
import json
import re
import zipfile
from datetime import datetime

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from crewai import Agent, Task, Crew, Process
from openai import OpenAI
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('dcf_pipeline')

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory job store
jobs = {}


# ═══════════════════════════════════════════════════════════════
#  STRUCTURED DATA EXTRACTION (post-agent JSON extraction)
# ═══════════════════════════════════════════════════════════════

EXTRACTION_PROMPT = """You are a financial data extraction specialist.
Given the complete DCF analysis output from 4 agents, extract ALL numerical
and textual data into the exact JSON structure below.

CRITICAL RULES:
- All monetary values must be numbers (not strings), in millions USD unless stated otherwise.
- All percentages must be plain numbers (e.g., 12.5 for 12.5%, NOT 0.125).
- If a value is truly missing, use null.
- The forecast array must have exactly 10 years (2026-2035).
- Do NOT invent data. Extract only what is present in the agent outputs.
- Return ONLY valid JSON, nothing else.

Required JSON structure:
{
  "company_name": "string",
  "ticker": "string or null",
  "country": "string",
  "industry": "string",
  "analysis_date": "YYYY-MM-DD",
  "method_summary": "string - 2-3 sentence professional DCF explanation",
  "assumptions": {
    "revenue_growth_rates": "string describing growth rate assumptions",
    "margin_assumptions": "string describing margin assumptions",
    "wacc": number,
    "terminal_growth_rate": number,
    "exit_multiple": number or null,
    "tax_rate": number,
    "risk_free_rate": number,
    "beta": number,
    "equity_risk_premium": number
  },
  "forecast": [
    {
      "year": 2026,
      "revenue": number,
      "revenue_growth_pct": number,
      "ebit_margin_pct": number,
      "ebit": number,
      "tax_rate": number,
      "nopat": number,
      "depreciation_amortization": number,
      "capex": number,
      "change_nwc": number,
      "fcff": number,
      "discount_factor": number,
      "pv_fcf": number
    }
  ],
  "terminal_value": number,
  "pv_terminal_value": number,
  "enterprise_value": number,
  "net_debt": number,
  "equity_value": number,
  "shares_outstanding": number,
  "intrinsic_value_per_share": number,
  "sensitivity": [
    {"wacc": number, "growth": number, "value_per_share": number}
  ],
  "risk_notes": ["string"],
  "validation_status": "Validated / Adjusted & Validated / Rejected",
  "validation_notes": ["string"]
}"""


def extract_structured_data(client, company_name, all_results):
    """Call OpenAI to extract structured JSON from the combined agent outputs."""
    combined = '\n\n'.join([
        f"=== AGENT {r['agent']}: {r['name']} ===\n{r['result']}"
        for r in all_results
    ])

    response = client.chat.completions.create(
        model='gpt-4.1-mini',
        messages=[
            {'role': 'system', 'content': EXTRACTION_PROMPT},
            {'role': 'user', 'content': f'Company analyzed: {company_name}\n\nFull agent outputs:\n{combined}'}
        ],
        temperature=0,
        max_tokens=8000,
        response_format={"type": "json_object"}
    )

    return json.loads(response.choices[0].message.content)


# ═══════════════════════════════════════════════════════════════
#  WORD GENERATION  (valuation_report.docx)
# ═══════════════════════════════════════════════════════════════

def _safe_num(val, fmt='{:,.1f}', fallback='N/A'):
    """Safely format a number, returning fallback if None."""
    if val is None:
        return fallback
    try:
        return fmt.format(float(val))
    except (ValueError, TypeError):
        return fallback


def create_word(data, company_name):
    """Create a professionally structured Word document valuation report."""
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # -- Title --
    title = doc.add_heading('DCF Valuation Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get('company_name', company_name))
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

    ticker = data.get('ticker')
    if ticker:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Ticker: {ticker} | {data.get('industry', 'N/A')} | {data.get('country', 'N/A')}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Date of Analysis: {data.get('analysis_date', datetime.now().strftime('%Y-%m-%d'))}")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_paragraph()  # spacer

    # -- 1. Summary of Method --
    doc.add_heading('1. Summary of Method', level=1)
    doc.add_paragraph(data.get('method_summary',
        'A Discounted Cash Flow (DCF) analysis was performed to estimate the intrinsic value '
        'of the company based on projected free cash flows discounted at the weighted average '
        'cost of capital (WACC).'))

    # -- 2. Key Assumptions --
    doc.add_heading('2. Key Assumptions', level=1)
    assumptions = data.get('assumptions', {})
    a_table = doc.add_table(rows=10, cols=2, style='Light Grid Accent 1')
    a_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_data = [
        ('Parameter', 'Value'),
        ('Revenue Growth Rates', str(assumptions.get('revenue_growth_rates', 'N/A'))),
        ('Margin Assumptions', str(assumptions.get('margin_assumptions', 'N/A'))),
        ('WACC', f"{_safe_num(assumptions.get('wacc'), '{:.2f}')}%"),
        ('Terminal Growth Rate', f"{_safe_num(assumptions.get('terminal_growth_rate'), '{:.2f}')}%"),
        ('Exit Multiple', _safe_num(assumptions.get('exit_multiple'), '{:.1f}x')),
        ('Tax Rate', f"{_safe_num(assumptions.get('tax_rate'), '{:.1f}')}%"),
        ('Risk-Free Rate', f"{_safe_num(assumptions.get('risk_free_rate'), '{:.2f}')}%"),
        ('Beta', _safe_num(assumptions.get('beta'), '{:.2f}')),
        ('Equity Risk Premium', f"{_safe_num(assumptions.get('equity_risk_premium'), '{:.2f}')}%"),
    ]
    for i, (label, value) in enumerate(headers_data):
        a_table.rows[i].cells[0].text = label
        a_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in a_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # -- 3. 10-Year Forecast Overview --
    doc.add_heading('3. 10-Year Forecast Overview', level=1)
    forecast = data.get('forecast', [])
    if forecast:
        fc_table = doc.add_table(rows=len(forecast) + 1, cols=5, style='Light Grid Accent 1')
        fc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fc_headers = ['Year', 'Revenue ($M)', 'EBIT ($M)', 'FCFF ($M)', 'PV of FCF ($M)']
        for j, h in enumerate(fc_headers):
            fc_table.rows[0].cells[j].text = h
            for paragraph in fc_table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for i, row in enumerate(forecast):
            fc_table.rows[i + 1].cells[0].text = str(row.get('year', ''))
            fc_table.rows[i + 1].cells[1].text = _safe_num(row.get('revenue'))
            fc_table.rows[i + 1].cells[2].text = _safe_num(row.get('ebit'))
            fc_table.rows[i + 1].cells[3].text = _safe_num(row.get('fcff'))
            fc_table.rows[i + 1].cells[4].text = _safe_num(row.get('pv_fcf'))
    else:
        doc.add_paragraph('Forecast data not available.')

    # -- 4. Valuation Summary --
    doc.add_heading('4. Valuation Summary', level=1)
    val_items = [
        ('Terminal Value ($M)', _safe_num(data.get('terminal_value'))),
        ('PV of Terminal Value ($M)', _safe_num(data.get('pv_terminal_value'))),
        ('Enterprise Value ($M)', _safe_num(data.get('enterprise_value'))),
        ('Net Debt ($M)', _safe_num(data.get('net_debt'))),
        ('Equity Value ($M)', _safe_num(data.get('equity_value'))),
        ('Shares Outstanding (M)', _safe_num(data.get('shares_outstanding'), '{:,.2f}')),
        ('Intrinsic Value Per Share', f"${_safe_num(data.get('intrinsic_value_per_share'), '{:,.2f}')}"),
    ]
    v_table = doc.add_table(rows=len(val_items) + 1, cols=2, style='Light Grid Accent 1')
    v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_table.rows[0].cells[0].text = 'Metric'
    v_table.rows[0].cells[1].text = 'Value'
    for paragraph in v_table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in v_table.rows[0].cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for i, (label, value) in enumerate(val_items):
        v_table.rows[i + 1].cells[0].text = label
        v_table.rows[i + 1].cells[1].text = value
        if i == len(val_items) - 1:
            for paragraph in v_table.rows[i + 1].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in v_table.rows[i + 1].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # -- 5. Sensitivity Analysis --
    doc.add_heading('5. Sensitivity Analysis', level=1)
    sensitivity = data.get('sensitivity', [])
    if sensitivity:
        s_table = doc.add_table(rows=len(sensitivity) + 1, cols=3, style='Light Grid Accent 1')
        s_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        s_headers = ['WACC (%)', 'Terminal Growth (%)', 'Value / Share ($)']
        for j, h in enumerate(s_headers):
            s_table.rows[0].cells[j].text = h
            for paragraph in s_table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for i, s in enumerate(sensitivity):
            s_table.rows[i + 1].cells[0].text = f"{_safe_num(s.get('wacc'), '{:.1f}')}%"
            s_table.rows[i + 1].cells[1].text = f"{_safe_num(s.get('growth'), '{:.1f}')}%"
            s_table.rows[i + 1].cells[2].text = f"${_safe_num(s.get('value_per_share'), '{:,.2f}')}"
    else:
        doc.add_paragraph('Sensitivity data not available.')

    # -- 6. Key Risk Notes --
    doc.add_heading('6. Key Risk Notes', level=1)
    risk_notes = data.get('risk_notes', [])
    if risk_notes:
        for note in risk_notes:
            doc.add_paragraph(note, style='List Bullet')
    else:
        doc.add_paragraph('No specific risk notes flagged.')

    # -- 7. Validation Status --
    doc.add_heading('7. Validation Status', level=1)
    v_status = data.get('validation_status', 'N/A')
    p_status = doc.add_paragraph()
    run_s = p_status.add_run(f'Status: {v_status}')
    run_s.bold = True
    if 'validated' in v_status.lower():
        run_s.font.color.rgb = RGBColor(0x48, 0xBB, 0x78)
    else:
        run_s.font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)

    val_notes = data.get('validation_notes', [])
    if val_notes:
        for note in val_notes:
            doc.add_paragraph(note, style='List Bullet')

    # -- Disclaimer --
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = disc.add_run(
        'This report was generated by DCF Production. '
        'It is intended for informational purposes only and does not constitute financial advice.'
    )
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════
#  EXCEL GENERATION  (dcf_10_year_forecast.xlsx — single sheet)
# ═══════════════════════════════════════════════════════════════

def create_excel(data):
    """Create a single-sheet DCF Excel workbook (DCF_10Y_Model)."""
    wb = Workbook()
    ws = wb.active
    ws.title = 'DCF_10Y_Model'

    # ── Styles ──
    hdr_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
    hdr_fill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
    hdr_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    data_font = Font(name='Calibri', size=10)
    data_align = Alignment(horizontal='right', vertical='center')
    lbl_font = Font(name='Calibri', bold=True, size=10)
    lbl_align = Alignment(horizontal='left', vertical='center')
    sum_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
    sum_font = Font(name='Calibri', bold=True, size=11, color='1B5E20')
    border = Border(
        left=Side(style='thin', color='D0D0D0'),
        right=Side(style='thin', color='D0D0D0'),
        top=Side(style='thin', color='D0D0D0'),
        bottom=Side(style='thin', color='D0D0D0')
    )
    NUM_FMT = '#,##0.0'
    PCT_FMT = '0.0%'
    DOLLAR_FMT = '$#,##0.00'

    # ── Column headers (row 1) ──
    headers = [
        'Year', 'Revenue ($M)', 'Revenue Growth %', 'EBIT Margin %',
        'EBIT ($M)', 'Tax Rate %', 'NOPAT ($M)', 'D&A ($M)',
        'Capex ($M)', 'Change in NWC ($M)', 'FCFF ($M)',
        'Discount Factor', 'PV of FCF ($M)'
    ]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hdr_font
        c.fill = hdr_fill
        c.alignment = hdr_align
        c.border = border

    # ── Forecast rows ──
    forecast = data.get('forecast', [])
    for i, row in enumerate(forecast):
        r = i + 2
        values = [
            row.get('year'),
            row.get('revenue'),
            (row.get('revenue_growth_pct') or 0) / 100.0,
            (row.get('ebit_margin_pct') or 0) / 100.0,
            row.get('ebit'),
            (row.get('tax_rate') or 0) / 100.0,
            row.get('nopat'),
            row.get('depreciation_amortization'),
            row.get('capex'),
            row.get('change_nwc'),
            row.get('fcff'),
            row.get('discount_factor'),
            row.get('pv_fcf'),
        ]
        for col, val in enumerate(values, 1):
            c = ws.cell(row=r, column=col, value=val)
            c.font = data_font
            c.alignment = data_align
            c.border = border
            if col == 1:
                c.number_format = '0'
                c.alignment = Alignment(horizontal='center')
            elif col in (3, 4, 6):
                c.number_format = PCT_FMT
            elif col == 12:
                c.number_format = '0.0000'
            else:
                c.number_format = NUM_FMT

    # ── Summary section (2 rows below last forecast) ──
    summary_start = len(forecast) + 3
    summary_items = [
        ('Terminal Value ($M)', data.get('terminal_value'), NUM_FMT),
        ('PV of Terminal Value ($M)', data.get('pv_terminal_value'), NUM_FMT),
        ('Enterprise Value ($M)', data.get('enterprise_value'), NUM_FMT),
        ('Net Debt ($M)', data.get('net_debt'), NUM_FMT),
        ('Equity Value ($M)', data.get('equity_value'), NUM_FMT),
        ('Shares Outstanding (M)', data.get('shares_outstanding'), '#,##0.00'),
        ('Intrinsic Value Per Share ($)', data.get('intrinsic_value_per_share'), DOLLAR_FMT),
    ]
    for idx, (label, value, fmt) in enumerate(summary_items):
        row = summary_start + idx
        lc = ws.cell(row=row, column=1, value=label)
        lc.font = lbl_font
        lc.alignment = lbl_align
        lc.border = border
        lc.fill = sum_fill

        vc = ws.cell(row=row, column=2, value=value)
        vc.font = sum_font if idx == len(summary_items) - 1 else data_font
        vc.alignment = data_align
        vc.border = border
        vc.fill = sum_fill
        vc.number_format = fmt

    # ── Column widths ──
    widths = [8, 16, 16, 14, 14, 12, 14, 12, 12, 16, 14, 14, 16]
    for i, w in enumerate(widths):
        ws.column_dimensions[chr(65 + i)].width = w

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


# ═══════════════════════════════════════════════════════════════
#  ZIP CREATION
# ═══════════════════════════════════════════════════════════════

def create_zip(word_bytes, excel_bytes):
    """Bundle valuation_report.docx + dcf_10_year_forecast.xlsx into a ZIP."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('valuation_report.docx', word_bytes)
        zf.writestr('dcf_10_year_forecast.xlsx', excel_bytes)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════
#  CREWAI  PIPELINE
# ═══════════════════════════════════════════════════════════════

def _check_cancelled(job_id: str) -> bool:
    """Return True if the job was cancelled and mark final state."""
    job = jobs.get(job_id)
    if not job:
        return True
    if job.get('cancelled'):
        logger.info('[Job %s] Job was cancelled by user. Stopping pipeline.', job_id[:8])
        job['status'] = 'cancelled'
        job['current_agent_name'] = 'Cancelled by user'
        return True
    return False


def run_dcf_pipeline(job_id, company_name, api_key, prompts):
    """Run the 4 CrewAI agents sequentially, then generate ZIP (Word + Excel)."""
    try:
        os.environ['OPENAI_API_KEY'] = api_key
        client = OpenAI(api_key=api_key)
        logger.info('[Job %s] === DCF PIPELINE STARTED for "%s" ===', job_id[:8], company_name)

        if _check_cancelled(job_id):
            return

        # ─── Agent 1: Company Existence Validation ───────────────────
        jobs[job_id]['current_agent'] = 1
        jobs[job_id]['current_agent_name'] = 'Company Existence Validation'
        logger.info('[Job %s] Agent 1 (Company Existence Validation) starting...', job_id[:8])

        agent1 = Agent(
            role='Company Existence Validator',
            goal=f'Verify if the company "{company_name}" exists and gather basic corporate info',
            backstory=prompts.get('agent1', 'You are a corporate intelligence verification agent.'),
            verbose=False,
            llm='gpt-4.1-mini'
        )
        task1 = Task(
            description=(
                f'Verify the existence and identity of the company: {company_name}. '
                f'Return structured information: Company Status [Exists / Does Not Exist / Uncertain], '
                f'exact legal name, ticker (if public), country, industry, official website, '
                f'and a 2-3 line factual description.\n\n'
                f'IMPORTANT RULES:\n'
                f'- For descriptive or marketing-style names (e.g., with slogans or long taglines), '
                f'assume the company may be real and search broadly (official website, business registries, '
                f'LinkedIn, credible directories, press releases).\n'
                f'- Use \"Does Not Exist\" ONLY when repeated, multi-source searching strongly indicates the entity '
                f'is fictional, a placeholder, or appears only in examples/training material.\n'
                f'- When evidence is sparse, ambiguous, or conflicting, prefer \"Uncertain\" and explain why.\n'
                f'- It is better to return \"Uncertain\" for a potentially real small/private company than to '
                f'incorrectly say \"Does Not Exist\".'
            ),
            agent=agent1,
            expected_output=(
                'Structured company verification report with status, legal name, '
                'ticker, country, industry, website, and description. Status must be clearly labeled as '
                '\"Company Status: [Exists/Does Not Exist/Uncertain]\".'
            )
        )
        crew1 = Crew(agents=[agent1], tasks=[task1], process=Process.sequential, verbose=False)
        result1 = crew1.kickoff()
        result1_str = str(result1)

        logger.info('[Job %s] Agent 1 completed. Result length: %d chars', job_id[:8], len(result1_str))

        # Try to parse a structured Company Status from Agent 1 output.
        status = None
        try:
            m = re.search(
                r'Company Status\\s*[:\\-]\\s*\\[?(?P<status>[A-Za-z \\-/]+)\\]?',
                result1_str,
                re.IGNORECASE,
            )
            if m:
                status = m.group('status').strip().lower()
                logger.info('[Job %s] Agent 1 parsed status: %s', job_id[:8], status)
        except Exception as parse_err:
            logger.warning('[Job %s] Failed to parse Company Status from Agent 1 output: %s', job_id[:8], parse_err)

        # Stop the pipeline when Agent 1 cannot confidently confirm existence.
        # If status is "Does Not Exist" OR "Uncertain", we do NOT continue to the next agents.
        if status in ('does not exist', 'nonexistent', 'does-not-exist', 'uncertain'):
            logger.warning(
                '[Job %s] Agent 1: Company status is %s. Stopping pipeline before next agents.',
                job_id[:8],
                status,
            )
            jobs[job_id]['agent_results'].append({
                'agent': 1, 'name': 'Company Existence Validation', 'result': result1_str
            })
            jobs[job_id]['status'] = 'error'
            safe_status = status or "Unknown"
            jobs[job_id]['error'] = (
                f'Company verification failed: The company "{company_name}" was marked as "{safe_status}" '
                'by the verification agent, so the DCF pipeline was stopped.'
            )
            return

        jobs[job_id]['agent_results'].append({
            'agent': 1, 'name': 'Company Existence Validation', 'result': result1_str
        })

        if _check_cancelled(job_id):
            return

        # ─── Agent 2: DCF Input Data Collection ─────────────────────
        jobs[job_id]['current_agent'] = 2
        jobs[job_id]['current_agent_name'] = 'DCF Input Data Collection'
        logger.info('[Job %s] Agent 2 (DCF Input Data Collection) starting...', job_id[:8])

        agent2 = Agent(
            role='Financial Data Collector',
            goal=f'Collect all required DCF input data for {company_name}',
            backstory=prompts.get('agent2', 'You are a senior financial analyst.'),
            verbose=False,
            llm='gpt-4.1-mini'
        )
        task2 = Task(
            description=(
                f'Using the verified company identity of {company_name}, collect all 5 critical '
                f'DCF input categories: '
                f'(1) 5+ years historical financials (revenue, EBITDA, EBIT, net income, FCF, growth, margins), '
                f'(2) balance sheet data (debt, cash, net debt, working capital, shares), '
                f'(3) reinvestment data (capex, D&A, WC trends), '
                f'(4) discount rate components (risk-free rate, beta, ERP, cost of debt, capital structure, WACC), '
                f'(5) terminal assumptions (GDP reference, industry growth, exit multiples). '
                f'Separate actual vs estimated data, list missing data and provide a data quality score.\n\n'
                f'Company verification info:\n{result1_str[:3000]}'
            ),
            agent=agent2,
            expected_output=(
                'Structured financial data with all 5 DCF input categories clearly separated, '
                'with data quality score.'
            )
        )
        crew2 = Crew(agents=[agent2], tasks=[task2], process=Process.sequential, verbose=False)
        result2 = crew2.kickoff()
        result2_str = str(result2)

        logger.info('[Job %s] Agent 2 completed. Result length: %d chars', job_id[:8], len(result2_str))
        jobs[job_id]['agent_results'].append({
            'agent': 2, 'name': 'DCF Input Data Collection', 'result': result2_str
        })

        if _check_cancelled(job_id):
            return

        # ─── Agent 3: DCF Calculation ────────────────────────────────
        jobs[job_id]['current_agent'] = 3
        jobs[job_id]['current_agent_name'] = 'DCF Calculation'
        logger.info('[Job %s] Agent 3 (DCF Calculation) starting...', job_id[:8])

        agent3 = Agent(
            role='Valuation Modeling Expert',
            goal=f'Build a complete 10-year DCF model for {company_name}',
            backstory=prompts.get('agent3', 'You are a valuation modeling expert.'),
            verbose=False,
            llm='gpt-4.1-mini'
        )
        task3 = Task(
            description=(
                f'Using the collected financial inputs for {company_name}, build a complete '
                f'10-year DCF model (2026-2035). Follow strict rules: '
                f'(1) revenue forecast via historical CAGR with base/conservative/optimistic scenarios, '
                f'(2) margin forecast via trend logic, '
                f'(3) FCFF = EBIT*(1-Tax)+D&A-Capex-delta_NWC, '
                f'(4) discount each year using WACC, '
                f'(5) compute terminal value using both perpetual growth and exit multiple, '
                f'(6) bridge EV to equity via net debt, '
                f'(7) compute intrinsic value per share. '
                f'Output structured text tables for forecast, FCF, PV, terminal value, EV, '
                f'equity value, per-share value plus sensitivity analysis (WACC +/-1%, g +/-0.5%) '
                f'and clearly labeled final scenario results.\n\n'
                f'Financial input data:\n{result2_str[:4000]}'
            ),
            agent=agent3,
            expected_output=(
                'Complete DCF model with forecast tables, FCF calculations, PV, terminal value, '
                'EV, equity value, per-share value, and sensitivity analysis.'
            )
        )
        crew3 = Crew(agents=[agent3], tasks=[task3], process=Process.sequential, verbose=False)
        result3 = crew3.kickoff()
        result3_str = str(result3)

        logger.info('[Job %s] Agent 3 completed. Result length: %d chars', job_id[:8], len(result3_str))
        jobs[job_id]['agent_results'].append({
            'agent': 3, 'name': 'DCF Calculation', 'result': result3_str
        })

        if _check_cancelled(job_id):
            return

        # ─── Agent 4: Validation & Realism Audit ────────────────────
        jobs[job_id]['current_agent'] = 4
        jobs[job_id]['current_agent_name'] = 'Validation & Realism Audit'
        logger.info('[Job %s] Agent 4 (Validation & Realism Audit) starting...', job_id[:8])

        agent4 = Agent(
            role='Financial Realism Auditor',
            goal=f'Audit and validate the DCF analysis for {company_name}',
            backstory=prompts.get('agent4', 'You are a financial realism audit agent.'),
            verbose=False,
            llm='gpt-4.1-mini'
        )
        task4 = Task(
            description=(
                f'Audit the DCF analysis for {company_name}. '
                f'Check growth realism vs GDP, margin consistency vs industry, reinvestment logic, '
                f'WACC/beta sanity, and terminal value dominance (<75% EV). '
                f'Independently recalculate key metrics, correct unrealistic assumptions if needed. '
                f'Return strictly structured machine-readable text formatted as: '
                f'Sheet 1 Company Summary, Sheet 2 Input Data, Sheet 3 Forecast Model, '
                f'Sheet 4 Valuation Summary (EV, Equity Value, Intrinsic Value/Share), '
                f'Sheet 5 Validation Notes (corrections, adjustments, risk flags), '
                f'with final status [Validated/Adjusted & Validated/Rejected].\n\n'
                f'Company verification:\n{result1_str[:1500]}\n\n'
                f'Financial data:\n{result2_str[:2500]}\n\n'
                f'DCF model:\n{result3_str[:4000]}'
            ),
            agent=agent4,
            expected_output=(
                'Structured audit report with 5 sections and final validation status '
                '[Validated / Adjusted & Validated / Rejected].'
            )
        )
        crew4 = Crew(agents=[agent4], tasks=[task4], process=Process.sequential, verbose=False)
        result4 = crew4.kickoff()
        result4_str = str(result4)

        logger.info('[Job %s] Agent 4 completed. Result length: %d chars', job_id[:8], len(result4_str))

        if 'rejected' in result4_str.lower():
            logger.warning('[Job %s] Agent 4: Analysis REJECTED. Stopping pipeline.', job_id[:8])
            jobs[job_id]['agent_results'].append({
                'agent': 4, 'name': 'Validation & Realism Audit', 'result': result4_str
            })
            jobs[job_id]['status'] = 'error'
            jobs[job_id]['error'] = (
                'Validation agent rejected the analysis. See agent 4 results for details.'
            )
            return

        jobs[job_id]['agent_results'].append({
            'agent': 4, 'name': 'Validation & Realism Audit', 'result': result4_str
        })

        if _check_cancelled(job_id):
            return

        # ─── Extract structured data ────────────────────────────────
        jobs[job_id]['current_agent'] = 0
        jobs[job_id]['current_agent_name'] = 'Extracting structured data...'
        logger.info('[Job %s] All 4 agents done. Extracting structured JSON data...', job_id[:8])

        structured = extract_structured_data(client, company_name, jobs[job_id]['agent_results'])

        # ─── Generate Word + Excel + ZIP ────────────────────────────
        jobs[job_id]['current_agent_name'] = 'Generating reports...'
        logger.info('[Job %s] Generating Word document and Excel file...', job_id[:8])

        word_bytes = create_word(structured, company_name)
        excel_bytes = create_excel(structured)
        zip_bytes = create_zip(word_bytes, excel_bytes)

        # Dynamic filename: companyname_valuation_YYYYMMDD.zip
        raw = company_name.split('-')[0].strip() if '-' in company_name else company_name
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', raw)
        safe_name = re.sub(r'_+', '_', safe_name).strip('_').lower()
        date_str = datetime.now().strftime('%Y%m%d')
        zip_filename = f'{safe_name}_valuation_{date_str}.zip'

        jobs[job_id]['zip_data'] = base64.b64encode(zip_bytes).decode('utf-8')
        jobs[job_id]['zip_filename'] = zip_filename
        jobs[job_id]['download_ready'] = True
        jobs[job_id]['status'] = 'complete'
        jobs[job_id]['current_agent'] = 0
        jobs[job_id]['current_agent_name'] = 'Complete'
        logger.info('[Job %s] === PIPELINE COMPLETE. ZIP ready: %s ===', job_id[:8], zip_filename)

    except Exception as e:
        logger.error('[Job %s] PIPELINE FAILED: %s', job_id[:8], str(e), exc_info=True)
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)


# ═══════════════════════════════════════════════════════════════
#  FASTAPI ROUTES
# ═══════════════════════════════════════════════════════════════

@app.post('/api/dcf/start')
def start_dcf(data: dict):
    """Start a DCF analysis pipeline."""
    if not data:
        raise HTTPException(status_code=400, detail='No data provided')

    company_name = data.get('company_name', '').strip()
    api_key = data.get('api_key', '').strip()
    prompts = data.get('prompts', {})

    if not company_name:
        raise HTTPException(status_code=400, detail='Company name is required')
    if not api_key or api_key == 'NO_KEY':
        raise HTTPException(status_code=400, detail='Please configure a valid OpenAI API key in Settings.')

    job_id = str(uuid.uuid4())
    logger.info('New DCF job started: %s for company "%s"', job_id[:8], company_name)
    jobs[job_id] = {
        'status': 'running',
        'current_agent': 1,
        'current_agent_name': 'Company Existence Validation',
        'agent_results': [],
        'error': None,
        'download_ready': False,
        'zip_data': None,
        'zip_filename': None,
        'company_name': company_name,
        'cancelled': False,
    }

    thread = threading.Thread(
        target=run_dcf_pipeline,
        args=(job_id, company_name, api_key, prompts)
    )
    thread.daemon = True
    thread.start()

    return {'job_id': job_id}


@app.get('/api/dcf/status/{job_id}')
def dcf_status(job_id: str):
    """Get the status of a DCF analysis job."""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail='Job not found')

    return {
        'status': job['status'],
        'current_agent': job['current_agent'],
        'current_agent_name': job['current_agent_name'],
        'agent_results': job['agent_results'],
        'error': job['error'],
        'download_ready': job['download_ready'],
        'zip_filename': job.get('zip_filename'),
        'cancelled': job.get('cancelled', False),
    }


@app.post('/api/dcf/cancel/{job_id}')
def dcf_cancel(job_id: str):
    """Request cancellation of a running DCF analysis job."""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail='Job not found')

    job['cancelled'] = True
    # Mark job as cancelled immediately from API point of view.
    if job.get('status') == 'running':
        job['status'] = 'cancelled'
        job['current_agent_name'] = 'Cancelled by user'

    logger.info('Cancellation requested for job %s', job_id[:8])
    return {
        'status': job['status'],
        'current_agent': job['current_agent'],
        'current_agent_name': job['current_agent_name'],
        'agent_results': job['agent_results'],
        'error': job['error'],
        'download_ready': job['download_ready'],
        'zip_filename': job.get('zip_filename'),
        'cancelled': True,
    }


@app.get('/api/dcf/download/{job_id}')
def dcf_download(job_id: str):
    """Download the ZIP file for a completed DCF analysis."""
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail='Job not found')

    if not job.get('download_ready') or not job.get('zip_data'):
        raise HTTPException(status_code=404, detail='Download is not ready yet')

    zip_bytes = base64.b64decode(job['zip_data'])
    filename = job.get('zip_filename', 'dcf_valuation.zip')

    return StreamingResponse(
        io.BytesIO(zip_bytes),
        media_type='application/zip',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.get('/api/health')
def health():
    return {'status': 'UP'}


if __name__ == '__main__':
    import uvicorn

    uvicorn.run(app, host='0.0.0.0', port=5000)
