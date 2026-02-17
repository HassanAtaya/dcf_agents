import io
from datetime import datetime
from typing import Any, Dict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import zipfile


# ═══════════════════════════════════════════════════════════════
#  WORD GENERATION  (valuation_report.docx)
# ═══════════════════════════════════════════════════════════════

def _safe_num(val, fmt='{:,.1f}', fallback='N/A'):
    """Safely format a number, returning fallback if None."""
    if val is None:
        return fallback
    try:
        return fmt.format(float(val))
    except (ValueError, TypeError):
        return fallback


def create_word(data: Dict[str, Any], company_name: str) -> bytes:
    """Create a professionally structured Word document valuation report."""
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # -- Title --
    title = doc.add_heading('DCF Valuation Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.get('company_name', company_name))
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    ticker = data.get('ticker')
    if ticker:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Ticker: {ticker} | {data.get('industry', 'N/A')} | {data.get('country', 'N/A')}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Date of Analysis: {data.get('analysis_date', datetime.now().strftime('%Y-%m-%d'))}")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_paragraph()  # spacer

    # -- 1. Summary of Method --
    doc.add_heading('1. Summary of Method', level=1)
    doc.add_paragraph(data.get('method_summary',
        'A Discounted Cash Flow (DCF) analysis was performed to estimate the intrinsic value '
        'of the company based on projected free cash flows discounted at the weighted average '
        'cost of capital (WACC).'))

    # -- 2. Key Assumptions --
    doc.add_heading('2. Key Assumptions', level=1)
    assumptions = data.get('assumptions', {})
    a_table = doc.add_table(rows=10, cols=2, style='Light Grid Accent 1')
    a_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_data = [
        ('Parameter', 'Value'),
        ('Revenue Growth Rates', str(assumptions.get('revenue_growth_rates', 'N/A'))),
        ('Margin Assumptions', str(assumptions.get('margin_assumptions', 'N/A'))),
        ('WACC', f"{_safe_num(assumptions.get('wacc'), '{:.2f}')}%"),
        ('Terminal Growth Rate', f"{_safe_num(assumptions.get('terminal_growth_rate'), '{:.2f}')}%"),
        ('Exit Multiple', _safe_num(assumptions.get('exit_multiple'), '{:.1f}x')),
        ('Tax Rate', f"{_safe_num(assumptions.get('tax_rate'), '{:.1f}')}%"),
        ('Risk-Free Rate', f"{_safe_num(assumptions.get('risk_free_rate'), '{:.2f}')}%"),
        ('Beta', _safe_num(assumptions.get('beta'), '{:.2f}')),
        ('Equity Risk Premium', f"{_safe_num(assumptions.get('equity_risk_premium'), '{:.2f}')}%"),
    ]
    for i, (label, value) in enumerate(headers_data):
        a_table.rows[i].cells[0].text = label
        a_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in a_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # -- 3. 10-Year Forecast Overview --
    doc.add_heading('3. 10-Year Forecast Overview', level=1)
    forecast = data.get('forecast', [])
    if forecast:
        fc_table = doc.add_table(rows=len(forecast) + 1, cols=5, style='Light Grid Accent 1')
        fc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fc_headers = ['Year', 'Revenue ($M)', 'EBIT ($M)', 'FCFF ($M)', 'PV of FCF ($M)']
        for j, h in enumerate(fc_headers):
            fc_table.rows[0].cells[j].text = h
            for paragraph in fc_table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for i, row in enumerate(forecast):
            fc_table.rows[i + 1].cells[0].text = str(row.get('year', ''))
            fc_table.rows[i + 1].cells[1].text = _safe_num(row.get('revenue'))
            fc_table.rows[i + 1].cells[2].text = _safe_num(row.get('ebit'))
            fc_table.rows[i + 1].cells[3].text = _safe_num(row.get('fcff'))
            fc_table.rows[i + 1].cells[4].text = _safe_num(row.get('pv_fcf'))
    else:
        doc.add_paragraph('Forecast data not available.')

    # -- 4. Valuation Summary --
    doc.add_heading('4. Valuation Summary', level=1)
    val_items = [
        ('Terminal Value ($M)', _safe_num(data.get('terminal_value'))),
        ('PV of Terminal Value ($M)', _safe_num(data.get('pv_terminal_value'))),
        ('Enterprise Value ($M)', _safe_num(data.get('enterprise_value'))),
        ('Net Debt ($M)', _safe_num(data.get('net_debt'))),
        ('Equity Value ($M)', _safe_num(data.get('equity_value'))),
        ('Shares Outstanding (M)', _safe_num(data.get('shares_outstanding'), '{:,.2f}')),
        ('Intrinsic Value Per Share', f"${_safe_num(data.get('intrinsic_value_per_share'), '{:,.2f}')}"),
    ]
    v_table = doc.add_table(rows=len(val_items) + 1, cols=2, style='Light Grid Accent 1')
    v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_table.rows[0].cells[0].text = 'Metric'
    v_table.rows[0].cells[1].text = 'Value'
    for paragraph in v_table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for paragraph in v_table.rows[0].cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True
    for i, (label, value) in enumerate(val_items):
        v_table.rows[i + 1].cells[0].text = label
        v_table.rows[i + 1].cells[1].text = value
        if i == len(val_items) - 1:
            for paragraph in v_table.rows[i + 1].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in v_table.rows[i + 1].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    # -- 5. Sensitivity Analysis --
    doc.add_heading('5. Sensitivity Analysis', level=1)
    sensitivity = data.get('sensitivity', [])
    if sensitivity:
        s_table = doc.add_table(rows=len(sensitivity) + 1, cols=3, style='Light Grid Accent 1')
        s_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        s_headers = ['WACC (%)', 'Terminal Growth (%)', 'Value / Share ($)']
        for j, h in enumerate(s_headers):
            s_table.rows[0].cells[j].text = h
            for paragraph in s_table.rows[0].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for i, s in enumerate(sensitivity):
            s_table.rows[i + 1].cells[0].text = f"{_safe_num(s.get('wacc'), '{:.1f}')}%"
            s_table.rows[i + 1].cells[1].text = f"{_safe_num(s.get('growth'), '{:.1f}')}%"
            s_table.rows[i + 1].cells[2].text = f"${_safe_num(s.get('value_per_share'), '{:,.2f}')}"
    else:
        doc.add_paragraph('Sensitivity data not available.')

    # -- 6. Key Risk Notes --
    doc.add_heading('6. Key Risk Notes', level=1)
    risk_notes = data.get('risk_notes', [])
    if risk_notes:
        for note in risk_notes:
            doc.add_paragraph(note, style='List Bullet')
    else:
        doc.add_paragraph('No specific risk notes flagged.')

    # -- 7. Validation Status --
    doc.add_heading('7. Validation Status', level=1)
    v_status = data.get('validation_status', 'N/A')
    p_status = doc.add_paragraph()
    run_s = p_status.add_run(f'Status: {v_status}')
    run_s.bold = True
    if 'validated' in v_status.lower():
        run_s.font.color.rgb = RGBColor(0x48, 0xBB, 0x78)
    else:
        run_s.font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)

    val_notes = data.get('validation_notes', [])
    if val_notes:
        for note in val_notes:
            doc.add_paragraph(note, style='List Bullet')

    # -- Disclaimer --
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = disc.add_run(
        'This report was generated by DCF Production. '
        'It is intended for informational purposes only and does not constitute financial advice.'
    )
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════
#  EXCEL GENERATION  (dcf_10_year_forecast.xlsx — single sheet)
# ═══════════════════════════════════════════════════════════════

def create_excel(data: Dict[str, Any]) -> bytes:
    """Create a single-sheet DCF Excel workbook (DCF_10Y_Model)."""
    wb = Workbook()
    ws = wb.active
    ws.title = 'DCF_10Y_Model'

    # ── Styles ──
    hdr_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
    hdr_fill = PatternFill(start_color='667EEA', end_color='667EEA', fill_type='solid')
    hdr_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    data_font = Font(name='Calibri', size=10)
    data_align = Alignment(horizontal='right', vertical='center')
    lbl_font = Font(name='Calibri', bold=True, size=10)
    lbl_align = Alignment(horizontal='left', vertical='center')
    sum_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
    sum_font = Font(name='Calibri', bold=True, size=11, color='1B5E20')
    border = Border(
        left=Side(style='thin', color='D0D0D0'),
        right=Side(style='thin', color='D0D0D0'),
        top=Side(style='thin', color='D0D0D0'),
        bottom=Side(style='thin', color='D0D0D0'),
    )
    NUM_FMT = '#,##0.0'
    PCT_FMT = '0.0%'
    DOLLAR_FMT = '$#,##0.00'

    # ── Column headers (row 1) ──
    headers = [
        'Year', 'Revenue ($M)', 'Revenue Growth %', 'EBIT Margin %',
        'EBIT ($M)', 'Tax Rate %', 'NOPAT ($M)', 'D&A ($M)',
        'Capex ($M)', 'Change in NWC ($M)', 'FCFF ($M)',
        'Discount Factor', 'PV of FCF ($M)',
    ]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hdr_font
        c.fill = hdr_fill
        c.alignment = hdr_align
        c.border = border

    # ── Forecast rows ──
    forecast = data.get('forecast', [])
    for i, row in enumerate(forecast):
        r = i + 2
        values = [
            row.get('year'),
            row.get('revenue'),
            (row.get('revenue_growth_pct') or 0) / 100.0,
            (row.get('ebit_margin_pct') or 0) / 100.0,
            row.get('ebit'),
            (row.get('tax_rate') or 0) / 100.0,
            row.get('nopat'),
            row.get('depreciation_amortization'),
            row.get('capex'),
            row.get('change_nwc'),
            row.get('fcff'),
            row.get('discount_factor'),
            row.get('pv_fcf'),
        ]
        for col, val in enumerate(values, 1):
            c = ws.cell(row=r, column=col, value=val)
            c.font = data_font
            c.alignment = data_align
            c.border = border
            if col == 1:
                c.number_format = '0'
                c.alignment = Alignment(horizontal='center')
            elif col in (3, 4, 6):
                c.number_format = PCT_FMT
            elif col == 12:
                c.number_format = '0.0000'
            else:
                c.number_format = NUM_FMT

    # ── Summary section (2 rows below last forecast) ──
    summary_start = len(forecast) + 3
    summary_items = [
        ('Terminal Value ($M)', data.get('terminal_value'), NUM_FMT),
        ('PV of Terminal Value ($M)', data.get('pv_terminal_value'), NUM_FMT),
        ('Enterprise Value ($M)', data.get('enterprise_value'), NUM_FMT),
        ('Net Debt ($M)', data.get('net_debt'), NUM_FMT),
        ('Equity Value ($M)', data.get('equity_value'), NUM_FMT),
        ('Shares Outstanding (M)', data.get('shares_outstanding'), '#,##0.00'),
        ('Intrinsic Value Per Share ($)', data.get('intrinsic_value_per_share'), DOLLAR_FMT),
    ]
    for idx, (label, value, fmt) in enumerate(summary_items):
        row = summary_start + idx
        lc = ws.cell(row=row, column=1, value=label)
        lc.font = lbl_font
        lc.alignment = lbl_align
        lc.border = border
        lc.fill = sum_fill

        vc = ws.cell(row=row, column=2, value=value)
        vc.font = sum_font if idx == len(summary_items) - 1 else data_font
        vc.alignment = data_align
        vc.border = border
        vc.fill = sum_fill
        vc.number_format = fmt

    # ── Column widths ──
    widths = [8, 16, 16, 14, 14, 12, 14, 12, 12, 16, 14, 14, 16]
    for i, w in enumerate(widths):
        ws.column_dimensions[chr(65 + i)].width = w

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


# ═══════════════════════════════════════════════════════════════
#  ZIP CREATION
# ═══════════════════════════════════════════════════════════════

def create_zip(word_bytes: bytes, excel_bytes: bytes) -> bytes:
    """Bundle valuation_report.docx + dcf_10_year_forecast.xlsx into a ZIP."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('valuation_report.docx', word_bytes)
        zf.writestr('dcf_10_year_forecast.xlsx', excel_bytes)
    buffer.seek(0)
    return buffer.getvalue()

